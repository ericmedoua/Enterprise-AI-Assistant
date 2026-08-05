from pathlib import Path

from docx import Document as WordDocument

from langchain_core.documents import Document


class DOCXExtractor:
    def extract(
        self,
        path: str,
    ) -> list[Document]:

        document = WordDocument(path)

        text = "\n".join(paragraph.text for paragraph in document.paragraphs)

        return [
            Document(
                page_content=text,
                metadata={
                    "source": Path(path).name,
                    "page": 1,
                },
            )
        ]
